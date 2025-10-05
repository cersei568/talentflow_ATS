# app.py
import streamlit as st
import fitz  # PyMuPDF
import re
import json
import os
import uuid
import sqlite3
from datetime import datetime
from collections import Counter
from io import BytesIO
import hashlib

import pandas as pd
import plotly.express as px
import plotly.graph_objects as go

# Try to import new OpenAI client first, fallback to old library if present
try:
    from openai import OpenAI as OpenAIClient
    OPENAI_NEW_AVAILABLE = True
except Exception:
    OPENAI_NEW_AVAILABLE = False
    try:
        import openai as old_openai
        OPENAI_OLD_AVAILABLE = True
    except Exception:
        OPENAI_OLD_AVAILABLE = False

# Try to import python-docx for styled offer generation
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False

# --------------------
# Config & Styling
# --------------------
st.set_page_config(
    page_title="Talentflow - ATS",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');

html, body, [class*="css"] { 
    font-family: 'Inter', sans-serif; 
}

.stApp { 
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    padding: 0;
}

.main-content {
    background: #ffffff;
    border-radius: 20px;
    padding: 2rem;
    margin: 1rem;
    box-shadow: 0 20px 60px rgba(0,0,0,0.3);
}

.candidate-card { 
    background: linear-gradient(135deg, #ffffff 0%, #f8f9fa 100%);
    padding: 1.5rem; 
    border-radius: 12px; 
    box-shadow: 0 4px 15px rgba(0,0,0,0.08); 
    margin-bottom: 1rem;
    border-left: 4px solid #667eea;
    transition: transform 0.2s, box-shadow 0.2s;
}

.candidate-card:hover {
    transform: translateY(-2px);
    box-shadow: 0 6px 20px rgba(0,0,0,0.12);
}

.metric-card {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    color: white;
    padding: 1.5rem;
    border-radius: 12px;
    text-align: center;
    box-shadow: 0 4px 15px rgba(102, 126, 234, 0.4);
}

.metric-value {
    font-size: 2.5rem;
    font-weight: 700;
    margin: 0.5rem 0;
}

.metric-label {
    font-size: 0.9rem;
    opacity: 0.9;
    text-transform: uppercase;
    letter-spacing: 1px;
}

.job-badge {
    display: inline-block;
    padding: 0.25rem 0.75rem;
    background: #667eea;
    color: white;
    border-radius: 20px;
    font-size: 0.85rem;
    margin: 0.25rem;
}

.skill-badge {
    display: inline-block;
    padding: 0.35rem 0.85rem;
    background: #e3f2fd;
    color: #1976d2;
    border-radius: 16px;
    font-size: 0.8rem;
    margin: 0.25rem;
    font-weight: 500;
}

.stage-badge {
    padding: 0.4rem 1rem;
    border-radius: 20px;
    font-size: 0.85rem;
    font-weight: 600;
    text-align: center;
}

.stage-received { background: #e3f2fd; color: #1976d2; }
.stage-screening { background: #fff3e0; color: #f57c00; }
.stage-interview { background: #f3e5f5; color: #7b1fa2; }
.stage-offer { background: #e8f5e9; color: #388e3c; }
.stage-hired { background: #c8e6c9; color: #2e7d32; }
.stage-rejected { background: #ffebee; color: #c62828; }

.score-excellent { color: #4caf50; font-weight: 700; }
.score-good { color: #8bc34a; font-weight: 600; }
.score-average { color: #ff9800; font-weight: 600; }
.score-poor { color: #f44336; font-weight: 600; }

h1, h2, h3 {
    color: #1a237e;
}

.stButton>button {
    border-radius: 8px;
    font-weight: 600;
    transition: all 0.3s;
}

.stButton>button:hover {
    transform: translateY(-2px);
    box-shadow: 0 4px 12px rgba(0,0,0,0.15);
}

.info-box {
    background: #e3f2fd;
    padding: 1rem;
    border-radius: 8px;
    border-left: 4px solid #1976d2;
    margin: 1rem 0;
}

.success-box {
    background: #e8f5e9;
    padding: 1rem;
    border-radius: 8px;
    border-left: 4px solid #4caf50;
    margin: 1rem 0;
}

.warning-box {
    background: #fff3e0;
    padding: 1rem;
    border-radius: 8px;
    border-left: 4px solid #ff9800;
    margin: 1rem 0;
}

.sidebar .sidebar-content {
    background: linear-gradient(180deg, #667eea 0%, #764ba2 100%);
}

.stTextInput>div>div>input, .stTextArea textarea {
    border-radius: 8px;
    border: 2px solid #e0e0e0;
    transition: border-color 0.3s;
}

.stTextInput>div>div>input:focus, .stTextArea textarea:focus {
    border-color: #667eea;
    box-shadow: 0 0 0 2px rgba(102, 126, 234, 0.1);
}

</style>
""", unsafe_allow_html=True)

# --------------------
# Constants & DB init
# --------------------
DB_PATH = "talentflow_v2.db"
SKILLS_DB = [
    "python", "java", "javascript", "typescript", "c++", "c#", "ruby", "go", "rust", "php",
    "django", "flask", "fastapi", "spring", "express", "react", "angular", "vue", "svelte",
    "node.js", "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch", "keras",
    "sql", "postgresql", "mysql", "mongodb", "redis", "elasticsearch", "cassandra",
    "aws", "azure", "gcp", "docker", "kubernetes", "jenkins", "gitlab", "github actions",
    "terraform", "ansible", "linux", "bash", "powershell", "rest api", "graphql", "grpc",
    "machine learning", "deep learning", "nlp", "computer vision", "data science",
    "spark", "hadoop", "kafka", "airflow", "excel", "tableau", "powerbi", "looker",
    "project management", "agile", "scrum", "kanban", "jira", "confluence",
    "git", "html", "css", "sass", "tailwind", "bootstrap", "webpack", "vite"
]

STAGE_COLORS = {
    "Received": "stage-received",
    "Screening": "stage-screening",
    "Interview": "stage-interview",
    "Offer": "stage-offer",
    "Hired": "stage-hired",
    "Rejected": "stage-rejected"
}

def init_db(path=DB_PATH):
    """Initialize SQLite database with proper schema"""
    try:
        # Ensure directory exists
        db_dir = os.path.dirname(path)
        if db_dir and not os.path.exists(db_dir):
            os.makedirs(db_dir)
        
        conn = sqlite3.connect(path, check_same_thread=False)
        cur = conn.cursor()
        
        # Jobs table
        cur.execute("""
        CREATE TABLE IF NOT EXISTS jobs (
            id TEXT PRIMARY KEY,
            title TEXT NOT NULL,
            required_skills TEXT,
            min_experience INTEGER DEFAULT 0,
            max_experience INTEGER DEFAULT 50,
            description TEXT,
            department TEXT,
            location TEXT,
            employment_type TEXT DEFAULT 'Full-time',
            salary_range TEXT,
            status TEXT DEFAULT 'Open',
            created_at TEXT,
            updated_at TEXT
        )""")
        
        # Candidates table
        cur.execute("""
        CREATE TABLE IF NOT EXISTS candidates (
            id TEXT PRIMARY KEY,
            job_id TEXT,
            filename TEXT,
            name TEXT,
            email TEXT,
            phone TEXT,
            skills TEXT,
            years_experience INTEGER DEFAULT 0,
            score INTEGER DEFAULT 0,
            explanation TEXT,
            stage TEXT DEFAULT 'Received',
            source TEXT,
            location TEXT,
            current_company TEXT,
            current_title TEXT,
            education TEXT,
            uploaded_at TEXT,
            updated_at TEXT,
            text TEXT,
            resume_hash TEXT,
            FOREIGN KEY (job_id) REFERENCES jobs(id)
        )""")
        
        # Evaluations table
        cur.execute("""
        CREATE TABLE IF NOT EXISTS evaluations (
            id TEXT PRIMARY KEY,
            candidate_id TEXT,
            evaluator TEXT,
            technical INTEGER,
            communication INTEGER,
            cultural INTEGER,
            problem_solving INTEGER DEFAULT 5,
            avg_score REAL,
            notes TEXT,
            recommendation TEXT,
            created_at TEXT,
            FOREIGN KEY (candidate_id) REFERENCES candidates(id)
        )""")
        
        # Activity log table
        cur.execute("""
        CREATE TABLE IF NOT EXISTS activity_log (
            id TEXT PRIMARY KEY,
            entity_type TEXT,
            entity_id TEXT,
            action TEXT,
            details TEXT,
            user TEXT,
            created_at TEXT
        )""")
        
        # Create indexes
        cur.execute("CREATE INDEX IF NOT EXISTS idx_candidates_job ON candidates(job_id)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_candidates_stage ON candidates(stage)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_evaluations_candidate ON evaluations(candidate_id)")
        
        conn.commit()
        return conn
        
    except Exception as e:
        st.error(f"❌ Database initialization failed: {e}")
        return None

def migrate_db(conn):
    """Migrate existing database to new schema"""
    if not conn:
        return
    
    cur = conn.cursor()
    
    try:
        # Check and add missing columns to jobs table
        cur.execute("PRAGMA table_info(jobs)")
        jobs_columns = [row[1] for row in cur.fetchall()]
        
        migrations_jobs = [
            ('status', "ALTER TABLE jobs ADD COLUMN status TEXT DEFAULT 'Open'"),
            ('max_experience', "ALTER TABLE jobs ADD COLUMN max_experience INTEGER DEFAULT 50"),
            ('department', "ALTER TABLE jobs ADD COLUMN department TEXT"),
            ('location', "ALTER TABLE jobs ADD COLUMN location TEXT"),
            ('employment_type', "ALTER TABLE jobs ADD COLUMN employment_type TEXT DEFAULT 'Full-time'"),
            ('salary_range', "ALTER TABLE jobs ADD COLUMN salary_range TEXT"),
            ('updated_at', "ALTER TABLE jobs ADD COLUMN updated_at TEXT"),
        ]
        
        for column, sql in migrations_jobs:
            if column not in jobs_columns:
                cur.execute(sql)
        
        # Check and add missing columns to candidates table
        cur.execute("PRAGMA table_info(candidates)")
        candidates_columns = [row[1] for row in cur.fetchall()]
        
        migrations_candidates = [
            ('location', "ALTER TABLE candidates ADD COLUMN location TEXT"),
            ('current_company', "ALTER TABLE candidates ADD COLUMN current_company TEXT"),
            ('current_title', "ALTER TABLE candidates ADD COLUMN current_title TEXT"),
            ('education', "ALTER TABLE candidates ADD COLUMN education TEXT"),
            ('updated_at', "ALTER TABLE candidates ADD COLUMN updated_at TEXT"),
            ('resume_hash', "ALTER TABLE candidates ADD COLUMN resume_hash TEXT"),
        ]
        
        for column, sql in migrations_candidates:
            if column not in candidates_columns:
                cur.execute(sql)
        
        # Check and add missing columns to evaluations table
        cur.execute("PRAGMA table_info(evaluations)")
        evals_columns = [row[1] for row in cur.fetchall()]
        
        migrations_evals = [
            ('problem_solving', "ALTER TABLE evaluations ADD COLUMN problem_solving INTEGER DEFAULT 5"),
            ('recommendation', "ALTER TABLE evaluations ADD COLUMN recommendation TEXT"),
        ]
        
        for column, sql in migrations_evals:
            if column not in evals_columns:
                cur.execute(sql)
        
        # Update existing records
        cur.execute("UPDATE jobs SET status='Open' WHERE status IS NULL")
        cur.execute("UPDATE jobs SET employment_type='Full-time' WHERE employment_type IS NULL")
        cur.execute("UPDATE jobs SET max_experience=50 WHERE max_experience IS NULL")
        cur.execute("UPDATE evaluations SET problem_solving=5 WHERE problem_solving IS NULL")
        
        conn.commit()
        
    except Exception as e:
        st.warning(f"⚠️ Migration warning: {e}")
        conn.rollback()

# Initialize Database Connection
try:
    conn = init_db()
    if conn:
        migrate_db(conn)
    else:
        st.error("❌ Failed to initialize database")
        st.info("💡 Please check file permissions and available disk space")
        st.stop()
except Exception as e:
    st.error(f"❌ Critical error: {e}")
    st.stop()

# --------------------
# Utility Functions
# --------------------
def calculate_file_hash(file_content):
    """Calculate hash of file content to detect duplicates"""
    return hashlib.md5(file_content).hexdigest()

def log_activity(entity_type, entity_id, action, details="", user="System"):
    """Log activity for audit trail"""
    if not conn:
        return
    try:
        cur = conn.cursor()
        cur.execute("""
        INSERT INTO activity_log (id, entity_type, entity_id, action, details, user, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (str(uuid.uuid4()), entity_type, entity_id, action, details, user, datetime.utcnow().isoformat()))
        conn.commit()
    except Exception as e:
        st.warning(f"Failed to log activity: {e}")

def extract_text_from_pdf_file(uploaded_file):
    """Extract text from PDF with better error handling"""
    try:
        data = uploaded_file.read()
        doc = fitz.open(stream=data, filetype="pdf")
        txt = "\n".join(page.get_text() for page in doc)
        doc.close()
        return txt.strip()
    except Exception as e:
        st.error(f"❌ PDF parsing error: {e}")
        return ""

def parse_contact_advanced(text):
    """Advanced contact information extraction"""
    # Email
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    email_match = re.search(email_pattern, text)
    
    # Phone
    phone_patterns = [
        r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
        r'(\+?\d{1,3}[-.\s]?)?\d{10}',
    ]
    phone = ""
    for pattern in phone_patterns:
        phone_match = re.search(pattern, text)
        if phone_match:
            phone = phone_match.group(0)
            break
    
    # Name (improved detection)
    name = ""
    lines = text.splitlines()
    for i, line in enumerate(lines[:10]):  # Check first 10 lines
        line = line.strip()
        if line and 2 <= len(line.split()) <= 4:
            if re.match(r'^[A-Z][a-z]+(?:\s[A-Z][a-z]+)*$', line):
                name = line
                break
    
    # Years of experience
    years_patterns = [
        r'(\d{1,2})\+?\s*(?:years?|yrs?)\s*(?:of\s*)?(?:experience|exp)',
        r'(?:experience|exp).*?(\d{1,2})\+?\s*(?:years?|yrs?)',
    ]
    years_exp = 0
    for pattern in years_patterns:
        years_match = re.search(pattern, text.lower())
        if years_match:
            years_exp = int(years_match.group(1))
            break
    
    # Location
    location = ""
    location_match = re.search(r'(?:Location|Address|Based in)[:\s]+([A-Za-z\s,]+)', text, re.IGNORECASE)
    if location_match:
        location = location_match.group(1).strip()
    
    # Current company
    current_company = ""
    company_match = re.search(r'(?:Currently at|Working at|@)\s+([A-Za-z0-9\s&]+)', text, re.IGNORECASE)
    if company_match:
        current_company = company_match.group(1).strip()
    
    return {
        "email": email_match.group(0) if email_match else "",
        "phone": phone,
        "name": name,
        "years_experience": years_exp,
        "location": location,
        "current_company": current_company
    }

def parse_skills_fallback(text):
    """Improved skills extraction"""
    t = text.lower()
    found = set()
    
    # Check for skills in database
    for skill in SKILLS_DB:
        # Use word boundaries for better matching
        pattern = r'\b' + re.escape(skill.lower()) + r'\b'
        if re.search(pattern, t):
            found.add(skill.title())
    
    return sorted(found)

def calculate_match_score(candidate_skills, required_skills, candidate_exp=0, min_exp=0):
    """Enhanced matching algorithm"""
    if not required_skills:
        return 75  # Default score if no requirements
    
    cand = set(s.lower() for s in candidate_skills)
    req = set(s.lower() for s in required_skills)
    
    matched = len(cand & req)
    skill_score = (matched / len(req)) * 70 if len(req) else 0
    
    # Experience bonus
    exp_score = 0
    if candidate_exp >= min_exp:
        exp_score = min(30, 30 * (candidate_exp / max(min_exp, 1)))
    
    total_score = int(round(skill_score + exp_score))
    return min(100, max(0, total_score))

def get_score_class(score):
    """Return CSS class based on score"""
    if score >= 80:
        return "score-excellent"
    elif score >= 60:
        return "score-good"
    elif score >= 40:
        return "score-average"
    else:
        return "score-poor"

def format_skills_html(skills):
    """Format skills as HTML badges"""
    if not skills:
        return "<span style='color: #999;'>No skills detected</span>"
    return " ".join([f"<span class='skill-badge'>{s}</span>" for s in skills[:10]])

# --------------------
# OpenAI helpers
# --------------------
def get_openai_client(api_key):
    """Initialize OpenAI client with proper error handling"""
    if not api_key or api_key.strip() == "":
        return None
    
    if OPENAI_NEW_AVAILABLE:
        try:
            return OpenAIClient(api_key=api_key.strip())
        except Exception as e:
            st.warning(f"⚠️ OpenAI initialization failed: {e}")
            return None
    
    if OPENAI_OLD_AVAILABLE:
        try:
            old_openai.api_key = api_key.strip()
            return old_openai
        except Exception:
            return None
    
    return None

def safe_extract_json(raw):
    """Robust JSON extraction from AI responses"""
    if not raw:
        return None
    
    if not isinstance(raw, str):
        raw = str(raw)
    
    # Try to find JSON object
    s = raw.find("{")
    e = raw.rfind("}")
    
    if s != -1 and e != -1 and e > s:
        candidate = raw[s:e+1]
        
        # Try direct parse
        try:
            return json.loads(candidate)
        except Exception:
            pass
        
        # Clean common issues
        candidate = re.sub(r',\s*}', '}', candidate)
        candidate = re.sub(r',\s*]', ']', candidate)
        candidate = re.sub(r'}\s*{', '},{', candidate)
        
        try:
            return json.loads(candidate)
        except Exception:
            pass
    
    # Try parsing as array
    try:
        return json.loads(raw)
    except Exception:
        return None

def ai_parse_and_score(client, cv_text, job_description, required_skills=None):
    """AI-powered CV parsing and scoring"""
    if not client:
        return None
    
    skills_hint = f"\nRequired skills to look for: {', '.join(required_skills)}" if required_skills else ""
    
    prompt = f"""You are an expert HR assistant analyzing a candidate's CV against a job description.

Extract the following information and return ONLY a valid JSON object:

{{
  "name": "Full name of candidate",
  "email": "Email address",
  "phone": "Phone number",
  "location": "Current location/city",
  "current_company": "Current employer",
  "current_title": "Current job title",
  "education": "Highest education degree",
  "skills": ["skill1", "skill2", ...],
  "years_experience": <number>,
  "score": <0-100>,
  "explanation": "Brief explanation of the match quality and key strengths/gaps"
}}

CV TEXT:
\"\"\"
{cv_text[:4000]}
\"\"\"

JOB DESCRIPTION:
\"\"\"
{job_description[:2000]}
\"\"\"
{skills_hint}

Return ONLY the JSON object, no additional text.
"""
    
    try:
        if OPENAI_NEW_AVAILABLE and isinstance(client, OpenAIClient):
            resp = client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=1000
            )
            raw = resp.choices[0].message.content
        else:
            resp = client.ChatCompletion.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.3,
                max_tokens=1000
            )
            raw = resp.choices[0].message.content
        
        parsed = safe_extract_json(raw)
        return parsed
        
    except Exception as e:
        st.warning(f"⚠️ AI parsing failed: {str(e)[:100]}")
        return None

# --------------------
# DB Operations
# --------------------
def save_job_db(job):
    """Save or update job in database"""
    if not conn:
        st.error("Database connection not available")
        return False
    
    try:
        cur = conn.cursor()
        job["updated_at"] = datetime.utcnow().isoformat()
        
        cur.execute("""
        INSERT OR REPLACE INTO jobs 
        (id, title, required_skills, min_experience, max_experience, description, 
         department, location, employment_type, salary_range, status, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            job["id"], job["title"], 
            json.dumps(job.get("required_skills", []), ensure_ascii=False),
            job.get("min_experience", 0), job.get("max_experience", 50),
            job.get("description", ""), job.get("department", ""),
            job.get("location", ""), job.get("employment_type", "Full-time"),
            job.get("salary_range", ""), job.get("status", "Open"),
            job.get("created_at", datetime.utcnow().isoformat()),
            job["updated_at"]
        ))
        conn.commit()
        log_activity("job", job["id"], "created" if "created_at" not in job else "updated", job["title"])
        return True
    except Exception as e:
        st.error(f"Error saving job: {e}")
        return False

def list_jobs_db(status=None):
    """List all jobs with optional status filter"""
    if not conn:
        st.error("Database connection not available")
        return []
    
    try:
        cur = conn.cursor()
        if status:
            cur.execute("SELECT * FROM jobs WHERE status=? ORDER BY created_at DESC", (status,))
        else:
            cur.execute("SELECT * FROM jobs ORDER BY created_at DESC")
        
        rows = cur.fetchall()
        cols = [d[0] for d in cur.description]
        
        jobs = []
        for r in rows:
            job = dict(zip(cols, r))
            try:
                job["required_skills"] = json.loads(job.get("required_skills", "[]"))
            except:
                job["required_skills"] = []
            jobs.append(job)
        
        return jobs
    except Exception as e:
        st.error(f"Error listing jobs: {e}")
        return []

def delete_job_db(job_id):
    """Delete a job and its candidates"""
    if not conn:
        return False
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM jobs WHERE id=?", (job_id,))
        cur.execute("DELETE FROM candidates WHERE job_id=?", (job_id,))
        conn.commit()
        log_activity("job", job_id, "deleted")
        return True
    except Exception as e:
        st.error(f"Error deleting job: {e}")
        return False

def save_candidate_db(candidate):
    """Save or update candidate in database"""
    if not conn:
        st.error("Database connection not available")
        return False
    
    try:
        cur = conn.cursor()
        candidate["updated_at"] = datetime.utcnow().isoformat()
        
        cur.execute("""
        INSERT OR REPLACE INTO candidates 
        (id, job_id, filename, name, email, phone, skills, years_experience, 
         score, explanation, stage, source, location, current_company, current_title,
         education, uploaded_at, updated_at, text, resume_hash)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            candidate["id"], candidate.get("job_id"), candidate.get("filename"),
            candidate.get("name"), candidate.get("email"), candidate.get("phone"),
            json.dumps(candidate.get("skills", []), ensure_ascii=False),
            candidate.get("years_experience", 0), candidate.get("score", 0),
            candidate.get("explanation", ""), candidate.get("stage", "Received"),
            candidate.get("source", ""), candidate.get("location", ""),
            candidate.get("current_company", ""), candidate.get("current_title", ""),
            candidate.get("education", ""),
            candidate.get("uploaded_at", datetime.utcnow().isoformat()),
            candidate["updated_at"], candidate.get("text", ""),
            candidate.get("resume_hash", "")
        ))
        conn.commit()
        return True
    except Exception as e:
        st.error(f"Error saving candidate: {e}")
        return False

def get_candidates_by_job(job_id=None, stage=None):
    """Get candidates with optional filters"""
    if not conn:
        return []
    
    try:
        cur = conn.cursor()
        query = "SELECT * FROM candidates"
        params = []
        
        conditions = []
        if job_id:
            conditions.append("job_id=?")
            params.append(job_id)
        if stage:
            conditions.append("stage=?")
            params.append(stage)
        
        if conditions:
            query += " WHERE " + " AND ".join(conditions)
        
        query += " ORDER BY score DESC, uploaded_at DESC"
        
        cur.execute(query, params)
        rows = cur.fetchall()
        cols = [d[0] for d in cur.description]
        
        candidates = []
        for r in rows:
            cand = dict(zip(cols, r))
            try:
                cand["skills"] = json.loads(cand.get("skills", "[]"))
            except:
                cand["skills"] = []
            candidates.append(cand)
        
        return candidates
    except Exception as e:
        st.error(f"Error fetching candidates: {e}")
        return []

def update_candidate_stage(candidate_id, new_stage):
    """Update candidate stage"""
    if not conn:
        return False
    try:
        cur = conn.cursor()
        cur.execute("UPDATE candidates SET stage=?, updated_at=? WHERE id=?", 
                    (new_stage, datetime.utcnow().isoformat(), candidate_id))
        conn.commit()
        log_activity("candidate", candidate_id, "stage_changed", new_stage)
        return True
    except Exception as e:
        st.error(f"Error updating stage: {e}")
        return False

def delete_candidate_db(candidate_id):
    """Delete a candidate"""
    if not conn:
        return False
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM candidates WHERE id=?", (candidate_id,))
        cur.execute("DELETE FROM evaluations WHERE candidate_id=?", (candidate_id,))
        conn.commit()
        log_activity("candidate", candidate_id, "deleted")
        return True
    except Exception as e:
        st.error(f"Error deleting candidate: {e}")
        return False

def save_evaluation_db(candidate_id, evaluator, technical, communication, cultural, problem_solving, notes="", recommendation=""):
    """Save interview evaluation"""
    if not conn:
        return 0
    
    try:
        avg = round((technical + communication + cultural + problem_solving) / 4.0, 2)
        cur = conn.cursor()
        rec_id = str(uuid.uuid4())
        
        cur.execute("""
        INSERT INTO evaluations 
        (id, candidate_id, evaluator, technical, communication, cultural, 
         problem_solving, avg_score, notes, recommendation, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            rec_id, candidate_id, evaluator, technical, communication, 
            cultural, problem_solving, avg, notes, recommendation,
            datetime.utcnow().isoformat()
        ))
        conn.commit()
        log_activity("evaluation", rec_id, "created", f"Score: {avg}")
        return avg
    except Exception as e:
        st.error(f"Error saving evaluation: {e}")
        return 0

def get_evaluations_by_candidate(candidate_id):
    """Get all evaluations for a candidate"""
    if not conn:
        return []
    try:
        cur = conn.cursor()
        cur.execute("SELECT * FROM evaluations WHERE candidate_id=? ORDER BY created_at DESC", (candidate_id,))
        rows = cur.fetchall()
        cols = [d[0] for d in cur.description]
        return [dict(zip(cols, r)) for r in rows]
    except Exception as e:
        st.error(f"Error fetching evaluations: {e}")
        return []

def get_all_evaluations():
    """Get all evaluations"""
    if not conn:
        return []
    try:
        cur = conn.cursor()
        cur.execute("SELECT * FROM evaluations ORDER BY created_at DESC")
        rows = cur.fetchall()
        cols = [d[0] for d in cur.description]
        return [dict(zip(cols, r)) for r in rows]
    except Exception as e:
        st.error(f"Error fetching all evaluations: {e}")
        return []

def check_duplicate_resume(file_hash):
    """Check if resume already exists"""
    if not conn:
        return None
    try:
        cur = conn.cursor()
        cur.execute("SELECT id, name, filename FROM candidates WHERE resume_hash=?", (file_hash,))
        return cur.fetchone()
    except Exception as e:
        st.warning(f"Error checking duplicates: {e}")
        return None

# --------------------
# Offer Letter Generation
# --------------------
def generate_offer_docx_bytes(candidate_info, job_info, offer_details):
    """Generate professional offer letter"""
    if not DOCX_AVAILABLE:
        return None
    
    try:
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Header
        header = doc.add_heading('EMPLOYMENT OFFER LETTER', 0)
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Date
        doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
        doc.add_paragraph()
        
        # Candidate info
        doc.add_paragraph(f"{candidate_info.get('name', 'Candidate')}")
        if candidate_info.get('email'):
            doc.add_paragraph(f"Email: {candidate_info['email']}")
        doc.add_paragraph()
        
        # Salutation
        doc.add_paragraph(f"Dear {candidate_info.get('name', 'Candidate')},")
        doc.add_paragraph()
        
        # Opening paragraph
        p = doc.add_paragraph()
        p.add_run("We are delighted to offer you the position of ").bold = False
        p.add_run(f"{job_info.get('title', 'Position')}").bold = True
        p.add_run(f" at {offer_details.get('company_name', 'Our Company')}. ")
        p.add_run("We believe your skills and experience will be a great asset to our team.")
        doc.add_paragraph()
        
        # Position Details
        doc.add_heading('Position Details', level=2)
        details_table = doc.add_table(rows=5, cols=2)
        details_table.style = 'Light Grid Accent 1'
        
        details = [
            ('Position:', job_info.get('title', '')),
            ('Department:', job_info.get('department', '')),
            ('Location:', job_info.get('location', offer_details.get('location', ''))),
            ('Employment Type:', job_info.get('employment_type', 'Full-time')),
            ('Start Date:', offer_details.get('start_date', 'TBD'))
        ]
        
        for i, (label, value) in enumerate(details):
            details_table.rows[i].cells[0].text = label
            details_table.rows[i].cells[1].text = str(value)
        
        doc.add_paragraph()
        
        # Compensation
        doc.add_heading('Compensation & Benefits', level=2)
        doc.add_paragraph(f"• Base Salary: ${offer_details.get('salary', 'TBD')} per annum")
        
        if offer_details.get('bonus'):
            doc.add_paragraph(f"• Performance Bonus: {offer_details['bonus']}")
        
        # Benefits
        if offer_details.get('benefits'):
            doc.add_paragraph("• Benefits Package:")
            for benefit in offer_details['benefits'].split('\n'):
                if benefit.strip():
                    doc.add_paragraph(f"  - {benefit.strip()}", style='List Bullet 2')
        
        doc.add_paragraph()
        
        # Closing
        doc.add_heading('Next Steps', level=2)
        doc.add_paragraph(
            "Please review this offer carefully. If you choose to accept, "
            "please sign and return this letter by "
            f"{offer_details.get('response_deadline', 'the specified date')}."
        )
        doc.add_paragraph()
        
        doc.add_paragraph(
            "We are excited about the possibility of you joining our team and "
            "look forward to your positive response."
        )
        doc.add_paragraph()
        
        # Signature lines
        doc.add_paragraph("Sincerely,")
        doc.add_paragraph()
        doc.add_paragraph("_" * 40)
        doc.add_paragraph(f"{offer_details.get('hiring_manager', 'Hiring Manager')}")
        doc.add_paragraph(f"{offer_details.get('company_name', 'Company Name')}")
        doc.add_paragraph()
        doc.add_paragraph()
        
        doc.add_paragraph("Acceptance:")
        doc.add_paragraph()
        doc.add_paragraph("_" * 40)
        doc.add_paragraph("Signature")
        doc.add_paragraph()
        doc.add_paragraph("_" * 40)
        doc.add_paragraph("Date")
        
        # Save to BytesIO
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        
        return bio
    except Exception as e:
        st.error(f"Error generating offer letter: {e}")
        return None

# --------------------
# Dashboard Metrics
# --------------------
def display_dashboard_metrics():
    """Display key metrics on dashboard"""
    if not conn:
        st.error("Database connection not available")
        return
        
    col1, col2, col3, col4, col5 = st.columns(5)
    
    # Total jobs
    try:
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM jobs WHERE status='Open'")
        open_jobs = cur.fetchone()[0]
    except Exception as e:
        st.error(f"Error fetching jobs: {e}")
        open_jobs = 0
    
    with col1:
        st.markdown(f"""
        <div class='metric-card'>
            <div class='metric-label'>Open Jobs</div>
            <div class='metric-value'>{open_jobs}</div>
        </div>
        """, unsafe_allow_html=True)
    
    # Total candidates
    try:
        cur.execute("SELECT COUNT(*) FROM candidates")
        total_candidates = cur.fetchone()[0]
    except Exception as e:
        st.error(f"Error fetching candidates: {e}")
        total_candidates = 0
    
    with col2:
        st.markdown(f"""
        <div class='metric-card'>
            <div class='metric-label'>Total Candidates</div>
            <div class='metric-value'>{total_candidates}</div>
        </div>
        """, unsafe_allow_html=True)
    
    # Active candidates (not hired/rejected)
    try:
        cur.execute("SELECT COUNT(*) FROM candidates WHERE stage NOT IN ('Hired', 'Rejected')")
        active_candidates = cur.fetchone()[0]
    except Exception as e:
        st.error(f"Error fetching active candidates: {e}")
        active_candidates = 0
    
    with col3:
        st.markdown(f"""
        <div class='metric-card'>
            <div class='metric-label'>Active Pipeline</div>
            <div class='metric-value'>{active_candidates}</div>
        </div>
        """, unsafe_allow_html=True)
    
    # In interview
    try:
        cur.execute("SELECT COUNT(*) FROM candidates WHERE stage='Interview'")
        in_interview = cur.fetchone()[0]
    except Exception as e:
        st.error(f"Error fetching interviews: {e}")
        in_interview = 0
    
    with col4:
        st.markdown(f"""
        <div class='metric-card'>
            <div class='metric-label'>In Interview</div>
            <div class='metric-value'>{in_interview}</div>
        </div>
        """, unsafe_allow_html=True)
    
    # Offers extended
    try:
        cur.execute("SELECT COUNT(*) FROM candidates WHERE stage IN ('Offer', 'Hired')")
        offers = cur.fetchone()[0]
    except Exception as e:
        st.error(f"Error fetching offers: {e}")
        offers = 0
    
    with col5:
        st.markdown(f"""
        <div class='metric-card'>
            <div class='metric-label'>Offers</div>
            <div class='metric-value'>{offers}</div>
        </div>
        """, unsafe_allow_html=True)

# --------------------
# Main App Layout
# --------------------
st.title("🎯 Talentflow")
st.markdown("### Professional Applicant Tracking System")
st.markdown("---")

# Sidebar
with st.sidebar:
    st.image("https://img.icons8.com/fluency/96/000000/best-employee.png", width=80)
    st.markdown("## ⚙️ Settings")
    
    # API Key
    default_key = ""
    try:
        default_key = st.secrets.get("openai", {}).get("api_key", "")
    except:
        pass
    
    api_key = st.text_input(
        "OpenAI API Key (Optional)", 
        type="password", 
        value=default_key,
        help="Enter your OpenAI API key for AI-powered resume parsing"
    )
    
    client = get_openai_client(api_key) if api_key else None
    
    if client:
        st.success("✅ AI Parsing Enabled")
    else:
        st.info("ℹ️ Using Rule-based Parsing")
    
    st.markdown("---")
    st.markdown("### 📊 Quick Stats")
    
    # Quick stats in sidebar
    try:
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM candidates WHERE stage='Received'")
        new_applications = cur.fetchone()[0]
        st.metric("New Applications", new_applications)
        
        cur.execute("SELECT COUNT(*) FROM candidates WHERE stage='Interview'")
        interviews = cur.fetchone()[0]
        st.metric("Scheduled Interviews", interviews)
    except:
        pass
    
    st.markdown("---")
    st.caption("Talentflow v2.0 | Powered by Streamlit")

# Dashboard metrics
display_dashboard_metrics()
st.markdown("---")

# Main tabs
tab_overview, tab_jobs, tab_candidates, tab_pipeline, tab_analytics, tab_offers = st.tabs([
    "📊 Overview", "💼 Jobs", "👥 Candidates", "🔄 Pipeline", "📈 Analytics", "📄 Offers"
])

# --------------------
# OVERVIEW TAB
# --------------------
with tab_overview:
    st.header("📊 Recruitment Overview")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.subheader("Recent Activity")
        
        # Get recent candidates
        recent_candidates = get_candidates_by_job()[:5]
        
        if recent_candidates:
            for cand in recent_candidates:
                score = cand.get('score', 0)
                score_class = get_score_class(score)
                stage_class = STAGE_COLORS.get(cand.get('stage', 'Received'), '')
                
                st.markdown(f"""
                <div class='candidate-card'>
                    <div style='display: flex; justify-content: space-between; align-items: center;'>
                        <div>
                            <h4 style='margin: 0;'>{cand.get('name') or cand.get('filename', 'Unknown')}</h4>
                            <p style='margin: 5px 0; color: #666;'>{cand.get('current_title', '')} 
                            {('at ' + cand.get('current_company', '')) if cand.get('current_company') else ''}</p>
                            <div style='margin-top: 10px;'>
                                {format_skills_html(cand.get('skills', []))}
                            </div>
                        </div>
                        <div style='text-align: right;'>
                            <div class='stage-badge {stage_class}'>{cand.get('stage', 'Received')}</div>
                            <p style='margin-top: 10px;' class='{score_class}'>Match: {score}%</p>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.info("No candidates yet. Start by uploading resumes!")
    
    with col2:
        st.subheader("Active Jobs")
        
        active_jobs = list_jobs_db(status="Open")
        
        if active_jobs:
            for job in active_jobs[:5]:
                # Count candidates for this job
                try:
                    cur = conn.cursor()
                    cur.execute("SELECT COUNT(*) FROM candidates WHERE job_id=?", (job['id'],))
                    candidate_count = cur.fetchone()[0]
                except:
                    candidate_count = 0
                
                st.markdown(f"""
                <div class='info-box'>
                    <h4 style='margin: 0;'>{job['title']}</h4>
                    <p style='margin: 5px 0;'>{candidate_count} candidate(s)</p>
                    <p style='margin: 5px 0; font-size: 0.9em; color: #666;'>
                        {job.get('department', '')} • {job.get('location', '')}
                    </p>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.info("No active jobs. Create one in the Jobs tab!")

# --------------------
# JOBS TAB
# --------------------
with tab_jobs:
    st.header("💼 Job Management")
    
    col1, col2 = st.columns([2, 1])
    
    with col1:
        with st.expander("➕ Create New Job", expanded=False):
            with st.form("create_job_form", clear_on_submit=True):
                j_title = st.text_input("Job Title*", placeholder="e.g., Senior Software Engineer")
                
                col_a, col_b = st.columns(2)
                with col_a:
                    j_department = st.text_input("Department", placeholder="e.g., Engineering")
                with col_b:
                    j_location = st.text_input("Location", placeholder="e.g., San Francisco, CA")
                
                j_employment = st.selectbox("Employment Type", 
                    ["Full-time", "Part-time", "Contract", "Internship"])
                
                col_c, col_d = st.columns(2)
                with col_c:
                    j_min_exp = st.number_input("Min. Years Experience", 0, 50, 2)
                with col_d:
                    j_max_exp = st.number_input("Max. Years Experience", 0, 50, 10)
                
                j_salary = st.text_input("Salary Range", placeholder="e.g., $80,000 - $120,000")
                
                j_skills = st.text_input("Required Skills (comma-separated)*", 
                    placeholder="e.g., Python, React, AWS")
                
                j_description = st.text_area("Job Description*", height=150,
                    placeholder="Enter detailed job description, responsibilities, and requirements...")
                
                submitted = st.form_submit_button("💾 Save Job", use_container_width=True)
                
                if submitted:
                    if not j_title or not j_skills or not j_description:
                        st.error("Please fill in all required fields (*)")
                    else:
                        job = {
                            "id": str(uuid.uuid4()),
                            "title": j_title.strip(),
                            "department": j_department.strip(),
                            "location": j_location.strip(),
                            "employment_type": j_employment,
                            "min_experience": int(j_min_exp),
                            "max_experience": int(j_max_exp),
                            "salary_range": j_salary.strip(),
                            "required_skills": [s.strip().title() for s in j_skills.split(",") if s.strip()],
                            "description": j_description.strip(),
                            "status": "Open",
                            "created_at": datetime.utcnow().isoformat()
                        }
                        if save_job_db(job):
                            st.success(f"✅ Job '{j_title}' created successfully!")
                            st.balloons()
        
        st.markdown("---")
        st.subheader("📋 All Jobs")
        
        jobs = list_jobs_db()
        
        if not jobs:
            st.info("No jobs created yet. Create your first job above!")
        else:
            # Filter options
            filter_col1, filter_col2 = st.columns(2)
            with filter_col1:
                status_filter = st.selectbox("Filter by Status", ["All", "Open", "Closed"])
            
            filtered_jobs = jobs if status_filter == "All" else [j for j in jobs if j.get('status') == status_filter]
            
            for job in filtered_jobs:
                # Count candidates
                try:
                    cur = conn.cursor()
                    cur.execute("SELECT COUNT(*) FROM candidates WHERE job_id=?", (job['id'],))
                    candidate_count = cur.fetchone()[0]
                except:
                    candidate_count = 0
                
                with st.expander(f"**{job['title']}** ({candidate_count} candidates) — {job.get('status', 'Open')}"):
                    col_info1, col_info2 = st.columns(2)
                    
                    with col_info1:
                        st.write("**Details:**")
                        st.write(f"🏢 Department: {job.get('department', 'N/A')}")
                        st.write(f"📍 Location: {job.get('location', 'N/A')}")
                        st.write(f"💼 Type: {job.get('employment_type', 'N/A')}")
                        st.write(f"💰 Salary: {job.get('salary_range', 'N/A')}")
                    
                    with col_info2:
                        st.write("**Requirements:**")
                        st.write(f"📅 Experience: {job.get('min_experience', 0)}-{job.get('max_experience', 50)} years")
                        st.write("**Skills:**")
                        st.markdown(format_skills_html(job.get('required_skills', [])), unsafe_allow_html=True)
                    
                    st.write("**Description:**")
                    st.write(job.get('description', 'No description provided'))
                    
                    # Actions
                    action_col1, action_col2, action_col3 = st.columns(3)
                    
                    with action_col1:
                        new_status = "Closed" if job.get('status') == "Open" else "Open"
                        if st.button(f"{'🔒 Close' if job.get('status') == 'Open' else '🔓 Reopen'}", 
                                   key=f"toggle_{job['id']}"):
                            try:
                                cur = conn.cursor()
                                cur.execute("UPDATE jobs SET status=?, updated_at=? WHERE id=?",
                                          (new_status, datetime.utcnow().isoformat(), job['id']))
                                conn.commit()
                                st.success(f"Job {new_status.lower()}")
                                st.rerun()
                            except Exception as e:
                                st.error(f"Error: {e}")
                    
                    with action_col2:
                        if st.button(f"🗑️ Delete", key=f"del_{job['id']}"):
                            if candidate_count > 0:
                                st.warning(f"Cannot delete job with {candidate_count} candidates. Remove candidates first.")
                            else:
                                if delete_job_db(job['id']):
                                    st.success("Job deleted")
                                    st.rerun()
    
    with col2:
        st.subheader("📊 Job Statistics")
        
        if jobs:
            total_jobs = len(jobs)
            open_jobs = len([j for j in jobs if j.get('status') == 'Open'])
            closed_jobs = total_jobs - open_jobs
            
            fig = go.Figure(data=[go.Pie(
                labels=['Open', 'Closed'],
                values=[open_jobs, closed_jobs],
                hole=.4,
                marker_colors=['#667eea', '#f56565']
            )])
            fig.update_layout(height=250, margin=dict(t=0, b=0, l=0, r=0))
            st.plotly_chart(fig, use_container_width=True)
            
            # Most in-demand skills
            all_req_skills = []
            for job in jobs:
                all_req_skills.extend(job.get('required_skills', []))
            
            if all_req_skills:
                skill_counts = Counter(all_req_skills)
                top_skills = skill_counts.most_common(10)
                
                st.write("**Most Required Skills:**")
                for skill, count in top_skills:
                    st.write(f"• {skill}: {count} job(s)")

# --------------------
# CANDIDATES TAB
# --------------------
with tab_candidates:
    st.header("👥 Candidate Management")
    
    # Job selector
    jobs = list_jobs_db(status="Open")
    job_options = {"All Jobs": None}
    job_options.update({j['title']: j['id'] for j in jobs})
    
    selected_job_title = st.selectbox("Filter by Job", list(job_options.keys()))
    selected_job_id = job_options[selected_job_title]
    selected_job = next((j for j in jobs if j['id'] == selected_job_id), None) if selected_job_id else None
    
    # Upload section
    with st.expander("📤 Upload New Resumes", expanded=True):
        col_up1, col_up2 = st.columns([3, 1])
        
        with col_up1:
            uploaded_files = st.file_uploader(
                "Select PDF resume files",
                type=["pdf"],
                accept_multiple_files=True,
                help="Upload one or more PDF resumes"
            )
        
        with col_up2:
            source = st.selectbox("Source", 
                ["LinkedIn", "Job Board", "Referral", "Career Fair", "Direct Application", "Other"])
            use_ai = st.checkbox("Use AI Parsing", value=bool(client))
        
        if uploaded_files:
            if not selected_job_id:
                st.warning("⚠️ Please select a job before uploading resumes")
            else:
                if st.button("🚀 Process Resumes", use_container_width=True, type="primary"):
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    added = 0
                    duplicates = 0
                    errors = 0
                    
                    for idx, file in enumerate(uploaded_files):
                        status_text.text(f"Processing {file.name}...")
                        
                        try:
                            # Check for duplicates
                            file_content = file.read()
                            file.seek(0)  # Reset file pointer
                            file_hash = calculate_file_hash(file_content)
                            
                            existing = check_duplicate_resume(file_hash)
                            if existing:
                                duplicates += 1
                                st.warning(f"⚠️ Duplicate resume detected: {file.name} (already uploaded as {existing[2]})")
                                continue
                            
                            # Extract text
                            text = extract_text_from_pdf_file(file)
                            
                            if not text:
                                errors += 1
                                st.error(f"❌ Could not extract text from {file.name}")
                                continue
                            
                            # Parse contact info
                            contact = parse_contact_advanced(text)
                            
                            # AI parsing
                            parsed = None
                            if use_ai and client:
                                with st.spinner(f"AI analyzing {file.name}..."):
                                    parsed = ai_parse_and_score(
                                        client, 
                                        text, 
                                        selected_job.get("description", ""),
                                        selected_job.get("required_skills", [])
                                    )
                            
                            # Fallback to rule-based
                            if not parsed:
                                skills = parse_skills_fallback(text)
                                parsed = {
                                    "email": contact.get("email", ""),
                                    "phone": contact.get("phone", ""),
                                    "location": contact.get("location", ""),
                                    "current_company": contact.get("current_company", ""),
                                    "current_title": "",
                                    "education": "",
                                    "skills": skills,
                                    "years_experience": contact.get("years_experience", 0),
                                    "score": calculate_match_score(
                                        skills, 
                                        selected_job.get("required_skills", []),
                                        contact.get("years_experience", 0),
                                        selected_job.get("min_experience", 0)
                                    ),
                                    "explanation": "Processed with rule-based parser"
                                }
                            
                            # Create candidate record
                            candidate = {
                                "id": str(uuid.uuid4()),
                                "job_id": selected_job_id,
                                "filename": file.name,
                                "name": parsed.get("name") or contact.get("name", ""),
                                "email": parsed.get("email", ""),
                                "phone": parsed.get("phone", ""),
                                "location": parsed.get("location", ""),
                                "current_company": parsed.get("current_company", ""),
                                "current_title": parsed.get("current_title", ""),
                                "education": parsed.get("education", ""),
                                "skills": parsed.get("skills", []),
                                "years_experience": int(parsed.get("years_experience", 0)),
                                "score": int(parsed.get("score", 0)),
                                "explanation": parsed.get("explanation", ""),
                                "stage": "Received",
                                "source": source,
                                "uploaded_at": datetime.utcnow().isoformat(),
                                "text": text,
                                "resume_hash": file_hash
                            }
                            
                            if save_candidate_db(candidate):
                                added += 1
                            
                        except Exception as e:
                            errors += 1
                            st.error(f"❌ Error processing {file.name}: {str(e)}")
                        
                        progress_bar.progress((idx + 1) / len(uploaded_files))
                    
                    status_text.empty()
                    progress_bar.empty()
                    
                    # Summary
                    if added > 0:
                        st.success(f"✅ Successfully added {added} candidate(s)")
                    if duplicates > 0:
                        st.info(f"ℹ️ Skipped {duplicates} duplicate(s)")
                    if errors > 0:
                        st.error(f"❌ {errors} error(s) occurred")
    
    # Manual candidate entry
    with st.expander("✍️ Add Candidate Manually"):
        with st.form("manual_candidate_form", clear_on_submit=True):
            m_col1, m_col2 = st.columns(2)
            
            with m_col1:
                m_name = st.text_input("Full Name*")
                m_email = st.text_input("Email*")
                m_phone = st.text_input("Phone")
            
            with m_col2:
                m_location = st.text_input("Location")
                m_company = st.text_input("Current Company")
                m_years = st.number_input("Years of Experience", 0, 50, 0)
            
            m_skills = st.text_input("Skills (comma-separated)*", 
                placeholder="e.g., Python, React, AWS")
            
            m_submitted = st.form_submit_button("➕ Add Candidate", use_container_width=True)
            
            if m_submitted:
                if not m_name or not m_email or not m_skills:
                    st.error("Please fill in required fields (*)")
                elif not selected_job_id:
                    st.error("Please select a job first")
                else:
                    skills_list = [s.strip().title() for s in m_skills.split(",") if s.strip()]
                    
                    candidate = {
                        "id": str(uuid.uuid4()),
                        "job_id": selected_job_id,
                        "filename": f"manual_{m_name.replace(' ', '_')}",
                        "name": m_name,
                        "email": m_email,
                        "phone": m_phone,
                        "location": m_location,
                        "current_company": m_company,
                        "skills": skills_list,
                        "years_experience": m_years,
                        "score": calculate_match_score(
                            skills_list,
                            selected_job.get("required_skills", []) if selected_job else [],
                            m_years,
                            selected_job.get("min_experience", 0) if selected_job else 0
                        ),
                        "explanation": "Manually added candidate",
                        "stage": "Received",
                        "source": source,
                        "uploaded_at": datetime.utcnow().isoformat()
                    }
                    
                    if save_candidate_db(candidate):
                        st.success(f"✅ Added {m_name} successfully!")
    
    st.markdown("---")
    
    # Filters
    col_filter1, col_filter2, col_filter3, col_filter4 = st.columns(4)
    
    with col_filter1:
        stage_filter = st.selectbox("Stage", ["All"] + ["Received", "Screening", "Interview", "Offer", "Hired", "Rejected"])
    
    with col_filter2:
        source_filter = st.selectbox("Source", ["All", "LinkedIn", "Job Board", "Referral", "Career Fair", "Direct Application", "Other"])
    
    with col_filter3:
        min_score = st.slider("Min Match Score", 0, 100, 0)
    
    with col_filter4:
        sort_by = st.selectbox("Sort By", ["Score (High to Low)", "Date (Recent)", "Name (A-Z)"])
    
    # Get candidates
    if stage_filter == "All":
        candidates = get_candidates_by_job(selected_job_id)
    else:
        candidates = get_candidates_by_job(selected_job_id, stage_filter)
    
    # Apply filters
    if source_filter != "All":
        candidates = [c for c in candidates if c.get('source') == source_filter]
    
    candidates = [c for c in candidates if c.get('score', 0) >= min_score]
    
    # Sort
    if sort_by == "Score (High to Low)":
        candidates.sort(key=lambda x: x.get('score', 0), reverse=True)
    elif sort_by == "Date (Recent)":
        candidates.sort(key=lambda x: x.get('uploaded_at', ''), reverse=True)
    elif sort_by == "Name (A-Z)":
        candidates.sort(key=lambda x: x.get('name', x.get('filename', '')).lower())
    
    # Display candidates
    st.subheader(f"📋 Candidates ({len(candidates)})")
    
    if not candidates:
        st.info("No candidates match your filters")
    else:
        for cand in candidates:
            score = cand.get('score', 0)
            score_class = get_score_class(score)
            stage = cand.get('stage', 'Received')
            stage_class = STAGE_COLORS.get(stage, '')
            
            with st.container():
                st.markdown(f"""
                <div class='candidate-card'>
                    <div style='display: flex; justify-content: space-between; align-items: start;'>
                        <div style='flex: 1;'>
                            <h3 style='margin: 0; color: #1a237e;'>{cand.get('name') or cand.get('filename', 'Unknown')}</h3>
                            <p style='margin: 5px 0; color: #666; font-size: 0.95em;'>
                                {cand.get('current_title', '')} 
                                {('at ' + cand.get('current_company', '')) if cand.get('current_company') else ''}
                            </p>
                            <p style='margin: 5px 0; color: #888; font-size: 0.9em;'>
                                📧 {cand.get('email', 'N/A')} | 
                                📱 {cand.get('phone', 'N/A')} | 
                                📍 {cand.get('location', 'N/A')}
                            </p>
                            <p style='margin: 5px 0; color: #888; font-size: 0.9em;'>
                                💼 {cand.get('years_experience', 0)} years exp | 
                                📅 {cand.get('source', 'Unknown')} | 
                                🕒 {cand.get('uploaded_at', '')[:10]}
                            </p>
                            <div style='margin-top: 10px;'>
                                {format_skills_html(cand.get('skills', []))}
                            </div>
                        </div>
                        <div style='text-align: right; min-width: 150px;'>
                            <div class='stage-badge {stage_class}' style='margin-bottom: 10px;'>{stage}</div>
                            <p class='{score_class}' style='font-size: 1.5em; margin: 5px 0;'>{score}%</p>
                            <p style='font-size: 0.85em; color: #666;'>Match Score</p>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                
                # Action buttons
                btn_col1, btn_col2, btn_col3, btn_col4 = st.columns(4)
                
                with btn_col1:
                    if st.button("📄 View Details", key=f"view_{cand['id']}", use_container_width=True):
                        st.session_state[f"show_detail_{cand['id']}"] = True
                
                with btn_col2:
                    if st.button("📝 Scorecard", key=f"score_{cand['id']}", use_container_width=True):
                        st.session_state[f"show_scorecard_{cand['id']}"] = True
                
                with btn_col3:
                    stages = ["Received", "Screening", "Interview", "Offer", "Hired", "Rejected"]
                    current_idx = stages.index(stage) if stage in stages else 0
                    
                    new_stage = st.selectbox(
                        "Stage",
                        stages,
                        index=current_idx,
                        key=f"stage_{cand['id']}",
                        label_visibility="collapsed"
                    )
                    
                    if new_stage != stage:
                        if update_candidate_stage(cand['id'], new_stage):
                            st.success(f"Moved to {new_stage}")
                            st.rerun()
                
                with btn_col4:
                    if st.button("🗑️", key=f"delete_{cand['id']}", use_container_width=True):
                        if delete_candidate_db(cand['id']):
                            st.success("Deleted")
                            st.rerun()
                
                # Detail modal
                if st.session_state.get(f"show_detail_{cand['id']}", False):
                    with st.expander("📄 Full Details", expanded=True):
                        detail_col1, detail_col2 = st.columns(2)
                        
                        with detail_col1:
                            st.write("**Personal Information:**")
                            st.write(f"Name: {cand.get('name', 'N/A')}")
                            st.write(f"Email: {cand.get('email', 'N/A')}")
                            st.write(f"Phone: {cand.get('phone', 'N/A')}")
                            st.write(f"Location: {cand.get('location', 'N/A')}")
                            st.write(f"Education: {cand.get('education', 'N/A')}")
                        
                        with detail_col2:
                            st.write("**Professional Information:**")
                            st.write(f"Current Company: {cand.get('current_company', 'N/A')}")
                            st.write(f"Current Title: {cand.get('current_title', 'N/A')}")
                            st.write(f"Experience: {cand.get('years_experience', 0)} years")
                            st.write(f"Source: {cand.get('source', 'N/A')}")
                        
                        st.write("**AI Explanation:**")
                        st.info(cand.get('explanation', 'No explanation available'))
                        
                        if cand.get('text'):
                            st.write("**Resume Text (Preview):**")
                            st.text_area("", value=cand['text'][:2000], height=200, key=f"text_{cand['id']}", disabled=True)
                        
                        if st.button("Close Details", key=f"close_detail_{cand['id']}"):
                            st.session_state[f"show_detail_{cand['id']}"] = False
                            st.rerun()
                
                # Scorecard modal
                if st.session_state.get(f"show_scorecard_{cand['id']}", False):
                    with st.expander("📝 Interview Scorecard", expanded=True):
                        with st.form(f"scorecard_form_{cand['id']}"):
                            sc_evaluator = st.text_input("Evaluator Name*", key=f"eval_name_{cand['id']}")
                            
                            sc_col1, sc_col2 = st.columns(2)
                            
                            with sc_col1:
                                sc_technical = st.slider("Technical Skills", 0, 10, 5, key=f"tech_{cand['id']}")
                                sc_communication = st.slider("Communication", 0, 10, 5, key=f"comm_{cand['id']}")
                            
                            with sc_col2:
                                sc_cultural = st.slider("Cultural Fit", 0, 10, 5, key=f"cult_{cand['id']}")
                                sc_problem = st.slider("Problem Solving", 0, 10, 5, key=f"prob_{cand['id']}")
                            
                            sc_notes = st.text_area("Notes", height=100, key=f"notes_{cand['id']}")
                            sc_recommendation = st.selectbox("Recommendation", 
                                ["Strong Hire", "Hire", "Maybe", "No Hire"], 
                                key=f"rec_{cand['id']}")
                            
                            sc_submit = st.form_submit_button("💾 Save Scorecard", use_container_width=True)
                            
                            if sc_submit:
                                if not sc_evaluator:
                                    st.error("Please enter evaluator name")
                                else:
                                    avg_score = save_evaluation_db(
                                        cand['id'], sc_evaluator, sc_technical, 
                                        sc_communication, sc_cultural, sc_problem,
                                        sc_notes, sc_recommendation
                                    )
                                    if avg_score > 0:
                                        st.success(f"✅ Scorecard saved! Average: {avg_score}/10")
                                        st.session_state[f"show_scorecard_{cand['id']}"] = False
                                        st.rerun()
                        
                        # Show existing evaluations
                        evals = get_evaluations_by_candidate(cand['id'])
                        if evals:
                            st.write("**Previous Evaluations:**")
                            for ev in evals:
                                st.markdown(f"""
                                <div class='info-box'>
                                    <strong>{ev['evaluator']}</strong> — Avg: {ev['avg_score']}/10
                                    <br>Tech: {ev['technical']} | Comm: {ev['communication']} | 
                                    Cultural: {ev['cultural']} | Problem: {ev['problem_solving']}
                                    <br><em>{ev.get('recommendation', '')}</em>
                                    <br>{ev.get('notes', '')}
                                </div>
                                """, unsafe_allow_html=True)
                        
                        if st.button("Close Scorecard", key=f"close_score_{cand['id']}"):
                            st.session_state[f"show_scorecard_{cand['id']}"] = False
                            st.rerun()
                
                st.markdown("---")

# --------------------
# PIPELINE TAB
# --------------------
with tab_pipeline:
    st.header("🔄 Recruitment Pipeline")
    
    # Get pipeline data
    stages = ["Received", "Screening", "Interview", "Offer", "Hired"]
    
    col1, col2, col3, col4, col5 = st.columns(5)
    cols = [col1, col2, col3, col4, col5]
    
    for idx, stage in enumerate(stages):
        candidates = get_candidates_by_job(stage=stage)
        
        with cols[idx]:
            st.markdown(f"""
            <div class='metric-card'>
                <div class='metric-label'>{stage}</div>
                <div class='metric-value'>{len(candidates)}</div>
            </div>
            """, unsafe_allow_html=True)
            
            for cand in candidates[:5]:  # Show top 5
                score = cand.get('score', 0)
                st.markdown(f"""
                <div style='background: white; padding: 8px; margin: 5px 0; border-radius: 6px; font-size: 0.85em;'>
                    <strong>{cand.get('name') or cand.get('filename', 'Unknown')[:20]}</strong>
                    <br><span style='color: {("#4caf50" if score >= 70 else "#ff9800" if score >= 50 else "#f44336")};'>
                    {score}%</span>
                </div>
                """, unsafe_allow_html=True)
            
            if len(candidates) > 5:
                st.caption(f"+{len(candidates) - 5} more")
    
    
    st.markdown("---")
    
    # Kanban board
    st.subheader("📊 Full Pipeline View")
    
    # Stage selector for moving candidates
    move_col1, move_col2, move_col3 = st.columns(3)
    
    with move_col1:
        all_candidates = get_candidates_by_job()
        cand_names = {f"{c.get('name') or c.get('filename', 'Unknown')} ({c.get('stage', 'Received')})": c['id'] 
                     for c in all_candidates}
        selected_cand_name = st.selectbox("Select Candidate to Move", [""] + list(cand_names.keys()))
    
    with move_col2:
        target_stage = st.selectbox("Move to Stage", stages + ["Rejected"])
    
    with move_col3:
        if st.button("🔄 Move Candidate", use_container_width=True, type="primary"):
            if selected_cand_name:
                cand_id = cand_names[selected_cand_name]
                update_candidate_stage(cand_id, target_stage)
                st.success(f"Moved to {target_stage}")
                st.rerun()
            else:
                st.warning("Please select a candidate")
    
    st.markdown("---")
    
    # Conversion funnel
    st.subheader("📉 Conversion Funnel")
    
    funnel_data = []
    for stage in stages:
        count = len(get_candidates_by_job(stage=stage))
        funnel_data.append({"Stage": stage, "Count": count})
    
    if funnel_data:
        fig = go.Figure(go.Funnel(
            y=[d['Stage'] for d in funnel_data],
            x=[d['Count'] for d in funnel_data],
            textinfo="value+percent initial",
            marker={"color": ["#e3f2fd", "#bbdefb", "#90caf9", "#64b5f6", "#42a5f5"]}
        ))
        fig.update_layout(height=400)
        st.plotly_chart(fig, use_container_width=True)
        
        # Conversion rates
        st.subheader("📊 Conversion Rates")
        
        if len(funnel_data) > 1:
            conv_col1, conv_col2, conv_col3, conv_col4 = st.columns(4)
            
            with conv_col1:
                if funnel_data[0]['Count'] > 0:
                    rate = (funnel_data[1]['Count'] / funnel_data[0]['Count']) * 100
                    st.metric("Received → Screening", f"{rate:.1f}%")
            
            with conv_col2:
                if funnel_data[1]['Count'] > 0:
                    rate = (funnel_data[2]['Count'] / funnel_data[1]['Count']) * 100
                    st.metric("Screening → Interview", f"{rate:.1f}%")
            
            with conv_col3:
                if funnel_data[2]['Count'] > 0:
                    rate = (funnel_data[3]['Count'] / funnel_data[2]['Count']) * 100
                    st.metric("Interview → Offer", f"{rate:.1f}%")
            
            with conv_col4:
                if funnel_data[3]['Count'] > 0:
                    rate = (funnel_data[4]['Count'] / funnel_data[3]['Count']) * 100
                    st.metric("Offer → Hired", f"{rate:.1f}%")

# --------------------
# ANALYTICS TAB
# --------------------
with tab_analytics:
    st.header("📈 Analytics & Insights")
    
    all_candidates = get_candidates_by_job()
    
    if not all_candidates:
        st.info("No data available yet. Start by adding candidates!")
    else:
        # Time series
        st.subheader("📅 Applications Over Time")
        
        df_timeline = pd.DataFrame([
            {"date": c['uploaded_at'][:10], "count": 1} 
            for c in all_candidates if c.get('uploaded_at')
        ])
        
        if not df_timeline.empty:
            df_timeline['date'] = pd.to_datetime(df_timeline['date'])
            df_grouped = df_timeline.groupby('date').count().reset_index()
            
            fig_time = px.line(df_grouped, x='date', y='count', 
                              title='Daily Applications',
                              markers=True)
            fig_time.update_layout(height=300)
            st.plotly_chart(fig_time, use_container_width=True)
        
        col_a1, col_a2 = st.columns(2)
        
        with col_a1:
            # Skills analysis
            st.subheader("🔧 Most Common Skills")
            
            all_skills = []
            for c in all_candidates:
                all_skills.extend(c.get('skills', []))
            
            if all_skills:
                skill_counts = Counter([s.title() for s in all_skills])
                top_skills = skill_counts.most_common(20)
                
                df_skills = pd.DataFrame(top_skills, columns=['Skill', 'Count'])
                
                fig_skills = px.bar(df_skills, x='Count', y='Skill', 
                                   orientation='h',
                                   title='Top 20 Skills in Candidate Pool',
                                   color='Count',
                                   color_continuous_scale='Blues')
                fig_skills.update_layout(height=500)
                st.plotly_chart(fig_skills, use_container_width=True)
        
        with col_a2:
            # Source analysis
            st.subheader("📊 Candidate Sources")
            
            source_counts = Counter([c.get('source', 'Unknown') for c in all_candidates])
            df_sources = pd.DataFrame(source_counts.items(), columns=['Source', 'Count'])
            
            fig_sources = px.pie(df_sources, values='Count', names='Source',
                                title='Applications by Source',
                                color_discrete_sequence=px.colors.qualitative.Set3)
            fig_sources.update_layout(height=400)
            st.plotly_chart(fig_sources, use_container_width=True)
        
        # Score distribution
        st.subheader("📊 Score Distribution")
        
        scores = [c.get('score', 0) for c in all_candidates]
        
        fig_scores = go.Figure()
        fig_scores.add_trace(go.Histogram(
            x=scores,
            nbinsx=20,
            marker_color='#667eea',
            opacity=0.7
        ))
        fig_scores.update_layout(
            title='Candidate Score Distribution',
            xaxis_title='Match Score (%)',
            yaxis_title='Number of Candidates',
            height=400
        )
        st.plotly_chart(fig_scores, use_container_width=True)
        
        # Top candidates
        st.subheader("🌟 Top Candidates")
        
        top_candidates = sorted(all_candidates, key=lambda x: x.get('score', 0), reverse=True)[:10]
        
        df_top = pd.DataFrame([
            {
                "Name": c.get('name') or c.get('filename', 'Unknown'),
                "Score": c.get('score', 0),
                "Experience": f"{c.get('years_experience', 0)} years",
                "Stage": c.get('stage', 'Received'),
                "Skills": len(c.get('skills', []))
            }
            for c in top_candidates
        ])
        
        st.dataframe(df_top, use_container_width=True, hide_index=True)
        
        # Experience distribution
        col_b1, col_b2 = st.columns(2)
        
        with col_b1:
            st.subheader("💼 Experience Distribution")
            
            exp_ranges = {
                "0-2 years": 0,
                "3-5 years": 0,
                "6-10 years": 0,
                "10+ years": 0
            }
            
            for c in all_candidates:
                exp = c.get('years_experience', 0)
                if exp <= 2:
                    exp_ranges["0-2 years"] += 1
                elif exp <= 5:
                    exp_ranges["3-5 years"] += 1
                elif exp <= 10:
                    exp_ranges["6-10 years"] += 1
                else:
                    exp_ranges["10+ years"] += 1
            
            df_exp = pd.DataFrame(exp_ranges.items(), columns=['Range', 'Count'])
            
            fig_exp = px.bar(df_exp, x='Range', y='Count',
                            title='Candidates by Experience Level',
                            color='Count',
                            color_continuous_scale='Viridis')
            fig_exp.update_layout(height=400)
            st.plotly_chart(fig_exp, use_container_width=True)
        
        with col_b2:
            # Stage distribution
            st.subheader("📍 Current Stage Distribution")
            
            stage_counts = Counter([c.get('stage', 'Received') for c in all_candidates])
            df_stages = pd.DataFrame(stage_counts.items(), columns=['Stage', 'Count'])
            
            fig_stages = px.bar(df_stages, x='Stage', y='Count',
                               title='Candidates by Stage',
                               color='Count',
                               color_continuous_scale='Blues')
            fig_stages.update_layout(height=400)
            st.plotly_chart(fig_stages, use_container_width=True)
        
        # Interview performance
        st.subheader("⭐ Interview Performance")
        
        all_evals = get_all_evaluations()
        
        if all_evals:
            df_evals = pd.DataFrame(all_evals)
            
            if len(df_evals) > 0:
                col_e1, col_e2, col_e3 = st.columns(3)
                
                with col_e1:
                    avg_technical = df_evals['technical'].mean()
                    st.metric("Avg Technical Score", f"{avg_technical:.1f}/10")
                
                with col_e2:
                    avg_comm = df_evals['communication'].mean()
                    st.metric("Avg Communication Score", f"{avg_comm:.1f}/10")
                
                with col_e3:
                    avg_cultural = df_evals['cultural'].mean()
                    st.metric("Avg Cultural Fit Score", f"{avg_cultural:.1f}/10")
                
                # Score trends over time
                df_evals['created_at'] = pd.to_datetime(df_evals['created_at'])
                df_evals_sorted = df_evals.sort_values('created_at')
                
                fig_eval_trend = go.Figure()
                
                fig_eval_trend.add_trace(go.Scatter(
                    x=df_evals_sorted['created_at'],
                    y=df_evals_sorted['technical'],
                    mode='lines+markers',
                    name='Technical',
                    line=dict(color='#667eea')
                ))
                
                fig_eval_trend.add_trace(go.Scatter(
                    x=df_evals_sorted['created_at'],
                    y=df_evals_sorted['communication'],
                    mode='lines+markers',
                    name='Communication',
                    line=dict(color='#f56565')
                ))
                
                fig_eval_trend.add_trace(go.Scatter(
                    x=df_evals_sorted['created_at'],
                    y=df_evals_sorted['cultural'],
                    mode='lines+markers',
                    name='Cultural',
                    line=dict(color='#48bb78')
                ))
                
                fig_eval_trend.update_layout(
                    title='Interview Scores Over Time',
                    xaxis_title='Date',
                    yaxis_title='Score (out of 10)',
                    height=400
                )
                
                st.plotly_chart(fig_eval_trend, use_container_width=True)
        else:
            st.info("No interview evaluations recorded yet")
        
        # Export data
        st.subheader("📥 Export Data")
        
        export_col1, export_col2 = st.columns(2)
        
        with export_col1:
            # Export candidates to CSV
            df_export = pd.DataFrame([
                {
                    "Name": c.get('name', ''),
                    "Email": c.get('email', ''),
                    "Phone": c.get('phone', ''),
                    "Score": c.get('score', 0),
                    "Stage": c.get('stage', ''),
                    "Experience": c.get('years_experience', 0),
                    "Skills": ", ".join(c.get('skills', [])),
                    "Source": c.get('source', ''),
                    "Uploaded": c.get('uploaded_at', '')
                }
                for c in all_candidates
            ])
            
            csv = df_export.to_csv(index=False).encode('utf-8')
            
            st.download_button(
                label="📥 Download Candidates CSV",
                data=csv,
                file_name=f"candidates_{datetime.now().strftime('%Y%m%d')}.csv",
                mime="text/csv",
                use_container_width=True
            )
        
        with export_col2:
            # Export evaluations to CSV
            if all_evals:
                df_evals_export = pd.DataFrame(all_evals)
                csv_evals = df_evals_export.to_csv(index=False).encode('utf-8')
                
                st.download_button(
                    label="📥 Download Evaluations CSV",
                    data=csv_evals,
                    file_name=f"evaluations_{datetime.now().strftime('%Y%m%d')}.csv",
                    mime="text/csv",
                    use_container_width=True
                )

# --------------------
# OFFERS TAB
# --------------------
# --------------------
# OFFERS TAB
# --------------------
with tab_offers:
    st.header("📄 Offer Letters")
    
    # Get candidates in Offer stage
    offer_candidates = get_candidates_by_job(stage="Offer")
    
    if not offer_candidates:
        st.info("No candidates in 'Offer' stage. Move candidates to Offer stage from the Pipeline tab.")
    else:
        st.subheader("Generate Offer Letter")
        
        # Select candidate
        cand_options = {
            f"{c.get('name') or c.get('filename', 'Unknown')} ({c.get('score', 0)}%)": c 
            for c in offer_candidates
        }
        
        selected_cand_name = st.selectbox("Select Candidate", list(cand_options.keys()))
        selected_candidate = cand_options[selected_cand_name]
        
        # Get job details
        if selected_candidate.get('job_id'):
            jobs = list_jobs_db()
            job = next((j for j in jobs if j['id'] == selected_candidate['job_id']), None)
        else:
            job = None
        
        # Offer details form
        with st.form("offer_letter_form"):
            st.markdown("### Offer Details")
            
            col_o1, col_o2 = st.columns(2)
            
            with col_o1:
                company_name = st.text_input("Company Name*", value="Talentflow Inc.")
                position = st.text_input("Position*", value=job['title'] if job else "")
                department = st.text_input("Department", value=job.get('department', '') if job else "")
                location = st.text_input("Work Location", value=job.get('location', '') if job else "")
            
            with col_o2:
                salary = st.text_input("Annual Salary*", value="80000", placeholder="e.g., 80000")
                bonus = st.text_input("Bonus/Equity (Optional)", placeholder="e.g., 10% annual bonus")
                start_date = st.date_input("Start Date*")
                response_deadline = st.date_input("Response Deadline")
            
            benefits = st.text_area(
                "Benefits Package (one per line)*",
                value="Health, Dental, and Vision Insurance\n401(k) with company match\nPaid Time Off (20 days)\nRemote work flexibility\nProfessional development budget",
                height=150
            )
            
            hiring_manager = st.text_input("Hiring Manager Name*", value="HR Director")
            
            generate_offer = st.form_submit_button("📄 Generate Offer Letter", use_container_width=True, type="primary")
        
        # Handle form submission (outside the form)
        if generate_offer:
            if not company_name or not position or not salary or not benefits or not hiring_manager:
                st.error("Please fill in all required fields (*)")
            else:
                if not DOCX_AVAILABLE:
                    st.error("❌ python-docx library is not installed. Please install it: `pip install python-docx`")
                else:
                    with st.spinner("Generating offer letter..."):
                        candidate_info = {
                            'name': selected_candidate.get('name') or selected_candidate.get('filename', 'Candidate'),
                            'email': selected_candidate.get('email', '')
                        }
                        
                        job_info = {
                            'title': position,
                            'department': department,
                            'location': location,
                            'employment_type': job.get('employment_type', 'Full-time') if job else 'Full-time'
                        }
                        
                        offer_details = {
                            'company_name': company_name,
                            'salary': salary,
                            'bonus': bonus,
                            'start_date': start_date.strftime('%B %d, %Y'),
                            'response_deadline': response_deadline.strftime('%B %d, %Y'),
                            'benefits': benefits,
                            'hiring_manager': hiring_manager,
                            'location': location
                        }
                        
                        docx_bytes = generate_offer_docx_bytes(candidate_info, job_info, offer_details)
                        
                        if docx_bytes:
                            # Store in session state to allow download outside form
                            st.session_state['generated_offer'] = {
                                'bytes': docx_bytes,
                                'filename': f"Offer_Letter_{candidate_info['name'].replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx",
                                'candidate_name': candidate_info['name'],
                                'position': position
                            }
                            
                            # Log activity
                            log_activity(
                                "offer", 
                                selected_candidate['id'], 
                                "generated", 
                                f"Offer letter for {position}",
                                hiring_manager
                            )
                            
                            st.success("✅ Offer letter generated successfully!")
                            st.balloons()
        
        # Download button (outside the form)
        if 'generated_offer' in st.session_state:
            offer_data = st.session_state['generated_offer']
            
            st.markdown("---")
            st.markdown("### 📥 Download Generated Offer")
            
            col_dl1, col_dl2 = st.columns([2, 1])
            
            with col_dl1:
                st.info(f"**Candidate:** {offer_data['candidate_name']}\n\n**Position:** {offer_data['position']}")
            
            with col_dl2:
                st.download_button(
                    label="📥 Download Offer Letter (DOCX)",
                    data=offer_data['bytes'],
                    file_name=offer_data['filename'],
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
                
                if st.button("Clear", use_container_width=True):
                    del st.session_state['generated_offer']
                    st.experimental_rerun()
        
        st.markdown("---")
        
        # Show all candidates in offer stage
        st.subheader("📋 All Candidates in Offer Stage")
        
        if not offer_candidates:
            st.info("No candidates currently in Offer stage")
        else:
            for cand in offer_candidates:
                with st.container():
                    st.markdown(f"""
                    <div class='success-box'>
                        <h4 style='margin: 0;'>{cand.get('name') or cand.get('filename', 'Unknown')}</h4>
                        <p style='margin: 5px 0;'>
                            📧 {cand.get('email', 'N/A')} | 
                            📱 {cand.get('phone', 'N/A')} | 
                            ⭐ Score: {cand.get('score', 0)}%
                        </p>
                        <p style='margin: 5px 0; font-size: 0.9em;'>
                            💼 {cand.get('years_experience', 0)} years experience | 
                            📅 Applied: {cand.get('uploaded_at', '')[:10]}
                        </p>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # Quick actions
                    action_col1, action_col2, action_col3 = st.columns(3)
                    
                    with action_col1:
                        if st.button(f"📧 Send Email", key=f"email_offer_{cand['id']}", use_container_width=True):
                            if cand.get('email'):
                                st.info(f"📧 Email: {cand['email']}")
                                st.code(f"mailto:{cand['email']}?subject=Job Offer - {job['title'] if job else 'Position'}")
                            else:
                                st.warning("No email address available")
                    
                    with action_col2:
                        if st.button(f"✅ Mark as Hired", key=f"hire_{cand['id']}", use_container_width=True):
                            update_candidate_stage(cand['id'], 'Hired')
                            st.success("Moved to Hired!")
                            st.experimental_rerun()
                    
                    with action_col3:
                        if st.button(f"❌ Withdraw Offer", key=f"withdraw_{cand['id']}", use_container_width=True):
                            update_candidate_stage(cand['id'], 'Rejected')
                            st.warning("Offer withdrawn")
                            st.experimental_rerun()
                    
                    st.markdown("---")
        
        # Bulk offer generation
        if len(offer_candidates) > 1:
            st.markdown("---")
            st.subheader("📦 Bulk Offer Generation")
            
            st.info("💡 Generate offer letters for multiple candidates at once")
            
            selected_for_bulk = st.multiselect(
                "Select candidates for bulk offer generation",
                [f"{c.get('name') or c.get('filename', 'Unknown')}" for c in offer_candidates],
                default=[]
            )
            
            if selected_for_bulk:
                st.write(f"Selected {len(selected_for_bulk)} candidate(s)")
                
                if st.button("🚀 Generate All Offers", type="primary", use_container_width=True):
                    st.info("Bulk generation feature coming soon! For now, please generate offers individually.")
# --------------------
# Sidebar Quick Actions
# --------------------
with st.sidebar:
    st.markdown("---")
    st.markdown("### ⚡ Quick Actions")
    
    # Quick interview scorecard
    if st.button("📝 Quick Scorecard", use_container_width=True):
        st.session_state['show_quick_scorecard'] = True
    
    if st.session_state.get('show_quick_scorecard', False):
        st.markdown("#### Interview Scorecard")
        
        # Get all candidates
        all_cands = get_candidates_by_job()
        active_cands = [c for c in all_cands if c.get('stage') not in ['Hired', 'Rejected']]
        
        if not active_cands:
            st.warning("No active candidates")
        else:
            quick_cand = st.selectbox(
                "Candidate",
                [f"{c.get('name') or c.get('filename', 'Unknown')}" for c in active_cands],
                key="quick_cand_select"
            )
            
            quick_evaluator = st.text_input("Your Name", key="quick_eval")
            
            quick_tech = st.slider("Technical", 0, 10, 5, key="quick_tech")
            quick_comm = st.slider("Communication", 0, 10, 5, key="quick_comm")
            quick_cult = st.slider("Cultural Fit", 0, 10, 5, key="quick_cult")
            quick_prob = st.slider("Problem Solving", 0, 10, 5, key="quick_prob")
            
            quick_notes = st.text_area("Quick Notes", key="quick_notes")
            quick_rec = st.selectbox("Recommendation", 
                ["Strong Hire", "Hire", "Maybe", "No Hire"],
                key="quick_rec")
            
            if st.button("💾 Submit", key="quick_submit", use_container_width=True):
                if not quick_evaluator:
                    st.error("Enter your name")
                else:
                    # Find candidate
                    cand = next((c for c in active_cands 
                               if (c.get('name') or c.get('filename', 'Unknown')) == quick_cand), None)
                    
                    if cand:
                        avg = save_evaluation_db(
                            cand['id'], quick_evaluator, quick_tech, 
                            quick_comm, quick_cult, quick_prob,
                            quick_notes, quick_rec
                        )
                        
                        st.success(f"✅ Saved! Avg: {avg:.1f}/10")
                        
                        # Auto-suggest next stage
                        if avg >= 7.5:
                            st.info("💡 Suggest: Move to Offer")
                        elif avg >= 5:
                            st.info("💡 Suggest: Schedule another interview")
                        else:
                            st.info("💡 Suggest: Consider rejection")
                        
                        st.session_state['show_quick_scorecard'] = False
                        st.rerun()
            
            if st.button("❌ Cancel", key="quick_cancel"):
                st.session_state['show_quick_scorecard'] = False
                st.rerun()

# --------------------
# Footer
# --------------------
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #666; padding: 2rem;'>
    <p><strong>Talentflow</strong> — Professional Applicant Tracking System</p>
    <p style='font-size: 0.85em;'>💡 Tip: Keep your data secure and backup regularly</p>
</div>
""", unsafe_allow_html=True)